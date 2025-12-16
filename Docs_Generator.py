import requests
from docx import Document

class Report:
    def __init__(self, url):
        self.url = url

    def get_api_data(self):
        try:
            r = requests.get(self.url)
            return r.json()
        except:
            print("API Error")
            return None

    def make_report(self, api_data, my_data):
        doc = Document()
        
        doc.add_heading("My Report", level=1)

        doc.add_heading("My Data", level=2)
        for k, v in my_data.items():
            doc.add_paragraph(f"{k}: {v}")

        doc.add_heading("API Data", level=2)
        for k, v in api_data.items():
            doc.add_paragraph(f"{k}: {v}")

        doc.save("sample report.docx")
        print("Report Created!")

url = "https://jsonplaceholder.typicode.com/posts/1"

obj = Report(url)

api = obj.get_api_data()

my_data = {
    "Name": "Akshat",
    "Branch": "CSE",
    "Project": "DOC Report Generator"
}

if api:
    obj.make_report(api, my_data)
